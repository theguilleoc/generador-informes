#!/usr/bin/env python3
"""
Generate INFORME .docx report from 5 DV Power CAT124A test .docx files.
Uses template.docx as base — preserves header/footer/formatting pixel-perfectly.
Replaces data values in tables and images from input files.
"""
import os, sys, shutil, zipfile, io, re, tempfile
from copy import deepcopy
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILES = {
    'O1':  os.path.join(BASE_DIR, 'O1.docx'),
    'C':   os.path.join(BASE_DIR, 'C.docx'),
    'OCO': os.path.join(BASE_DIR, 'O_03s_CO.docx'),
    'CO':  os.path.join(BASE_DIR, 'CO.docx'),
    'O2':  os.path.join(BASE_DIR, 'O2.docx'),
}
TEMPLATE_PATH = os.path.join(BASE_DIR, 'template.docx')
OUTPUT_PATH   = os.path.join(BASE_DIR, 'INFORME_GENERADO.docx')

W_NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
A_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
VML_NS = 'urn:schemas-microsoft-com:vml'

# =============================================================================
# PARSING INPUT .docx FILES
# =============================================================================

def extract_images_from_docx(filepath):
    """Extract images from an input .docx in document order (by rId reference)."""
    images = []
    with zipfile.ZipFile(filepath, 'r') as z:
        # Read rels
        rels_xml = z.read('word/_rels/document.xml.rels')
        rels_doc = etree.fromstring(rels_xml)
        rel_map = {}
        for rel in rels_doc:
            rel_map[rel.get('Id')] = rel.get('Target')

        # Read document.xml, find blip references in order
        doc_xml = z.read('word/document.xml')
        doc_tree = etree.fromstring(doc_xml)
        seen_rids = []
        for blip in doc_tree.iter(qn('a:blip')):
            rId = blip.get(qn('r:embed'))
            if rId and rId not in seen_rids:
                seen_rids.append(rId)

        # Extract images in document order
        for rId in seen_rids:
            target = rel_map.get(rId, '')
            if target:
                img_path = 'word/' + target if not target.startswith('word/') else target
                try:
                    images.append(z.read(img_path))
                except KeyError:
                    print(f"  Warning: image {img_path} not found in {filepath}")
    return images


def parse_table_rows(doc, table_idx):
    """Get rows of a table as lists of cell texts (handling merged cells)."""
    if table_idx >= len(doc.tables):
        return []
    tbl = doc.tables[table_idx]
    rows = []
    for row in tbl.rows:
        seen = set()
        cells = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen:
                seen.add(cid)
                cells.append(cell.text.strip())
        rows.append(cells)
    return rows


def is_page_break_row(row):
    return len(row) >= 4 and row[0] == 'Test lab'

def is_type_header_row(row):
    return len(row) >= 4 and row[0] == 'Type'

def is_empty_row(row):
    return all(c.strip() == '' for c in row)


def parse_data_sections(rows):
    """Parse data sections from report table rows (table 1 of input .docx)."""
    sections = []
    cur = None
    i = 15
    saw_break = False
    while i < len(rows):
        row = rows[i]
        if is_page_break_row(row):
            saw_break = True; i += 1; continue
        if is_type_header_row(row) or (len(row) == 1 and row[0] == ''):
            i += 1; continue
        if is_empty_row(row):
            i += 1; continue
        ne = [c for c in row if c.strip()]
        if len(ne) == 1 and len(row) <= 4:
            if saw_break and cur and cur['name'] == ne[0]:
                saw_break = False; i += 1
                if i < len(rows):
                    nr = [c for c in rows[i] if c.strip()]
                    if len(nr) > 1 and nr[0] in ('Designation', 'Name', 'Coil'):
                        i += 1
                continue
            cur = {'name': ne[0], 'columns': [], 'data': []}
            sections.append(cur)
            saw_break = False; i += 1; continue
        if len(ne) > 1 and ne[0] in ('Designation', 'Name', 'Coil'):
            if cur:
                cur['columns'] = ne
            i += 1; continue
        if cur and len(ne) >= 3:
            cl = [c.strip() for c in row]
            s = 0
            while s < len(cl) and cl[s] == '':
                s += 1
            d = cl[s:]
            while d and d[-1] == '':
                d.pop()
            if len(d) >= 3:
                cur['data'].append(d)
        i += 1
    return sections


def parse_summary_sections(rows):
    """Parse summary sections (table 0 of input .docx)."""
    sections = []
    cur = None
    for row in rows:
        if is_page_break_row(row) or is_type_header_row(row):
            continue
        if len(row) == 1 and row[0] == '':
            continue
        if is_empty_row(row):
            continue
        ne = [c for c in row if c.strip()]
        if len(ne) == 1 and len(row) == 2:
            cur = {'name': ne[0], 'columns': [], 'data': []}
            sections.append(cur)
            continue
        if cur and not cur['columns'] and len(ne) > 1 and ne[0] in ('Designation', 'Name', 'Coil'):
            cur['columns'] = ne
            continue
        if cur and len(ne) >= 3:
            cur['data'].append(ne)
    return sections


def parse_input_file(filepath, slot):
    """Parse an input .docx test report."""
    print(f"  Parsing {slot}: {os.path.basename(filepath)}")
    doc = Document(filepath)
    t0_rows = parse_table_rows(doc, 0)  # summary
    t1_rows = parse_table_rows(doc, 1)  # report

    header = {}
    if t1_rows:
        header['type'] = t1_rows[0][1] if len(t1_rows[0]) > 1 else ''
        header['serial'] = t1_rows[0][3] if len(t1_rows[0]) > 3 else ''

    params = {}
    device_settings = {}
    for i in range(4, min(10, len(t1_rows))):
        r = t1_rows[i]
        if len(r) >= 3 and r[1].strip():
            params[r[1].strip()] = r[2].strip()
        if len(r) >= 5 and r[3].strip():
            device_settings[r[3].strip()] = r[4].strip()

    sections = parse_data_sections(t1_rows)
    summary_sections = parse_summary_sections(t0_rows)
    images = extract_images_from_docx(filepath)

    return {
        'slot': slot,
        'header': header,
        'params': params,
        'device_settings': device_settings,
        'sections': sections,
        'summary_sections': summary_sections,
        'images': images,
    }


# =============================================================================
# TEMPLATE MODIFICATION
# =============================================================================

def set_cell_text(cell, text):
    """Replace cell text preserving formatting of first run."""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = str(text)
            for r in p.runs[1:]:
                r.text = ''
        else:
            p.text = str(text)


def center_all_content(doc):
    """Center all body tables and VML image paragraphs. Fix VML absolute positioning."""
    NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    VML_NS_URI = 'urn:schemas-microsoft-com:vml'

    # --- Center all tables ---
    for tbl in doc.tables:
        tblPr = tbl._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._tbl.insert(0, tblPr)
        for old in tblPr.findall(qn('w:jc')):
            tblPr.remove(old)
        for old in tblPr.findall(qn('w:tblInd')):
            tblPr.remove(old)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)

    body = doc.element.body

    # --- Fix VML shapes: remove absolute positioning, keep original dimensions ---
    for shape in body.iter('{' + VML_NS_URI + '}shape'):
        imgdata = shape.find('{' + VML_NS_URI + '}imagedata')
        if imgdata is None:
            continue
        style = shape.get('style', '')
        # Extract original width and height
        w_match = re.search(r'width:\s*[\d.]+(?:pt|in)', style)
        h_match = re.search(r'height:\s*[\d.]+(?:pt|in)', style)
        if w_match and h_match:
            shape.set('style',
                       f'{w_match.group()};{h_match.group()};visibility:visible')
        # Remove wrap element (only used with absolute positioning)
        for wrap in shape.findall('{' + VML_NS_URI + '}wrap'):
            shape.remove(wrap)

    # --- Center all paragraphs that contain images and set zero spacing ---
    for p in body.findall('.//w:p', NSMAP):
        if not p.findall('.//w:pict', NSMAP):
            continue
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        for old in pPr.findall(qn('w:jc')):
            pPr.remove(old)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        # Zero paragraph spacing so image pairs stay together
        for old in pPr.findall(qn('w:spacing')):
            pPr.remove(old)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        pPr.append(spacing)

    # --- Remove empty paragraphs between consecutive image paragraphs ---
    children = list(body)
    to_remove = []
    for i, child in enumerate(children):
        if child.tag != qn('w:p'):
            continue
        # Skip if paragraph has images or text
        if child.findall('.//' + qn('w:pict')) or child.findall('.//' + qn('w:t')):
            continue
        # Empty paragraph — check if it sits between two image paragraphs
        has_img_before = False
        for j in range(i - 1, -1, -1):
            prev = children[j]
            if prev.tag == qn('w:tbl'):
                break
            if prev.tag == qn('w:p'):
                if prev.findall('.//' + qn('w:pict')):
                    has_img_before = True
                    break
                if prev.findall('.//' + qn('w:t')):
                    break
        if not has_img_before:
            continue
        has_img_after = False
        for j in range(i + 1, len(children)):
            nxt = children[j]
            if nxt.tag == qn('w:tbl'):
                break
            if nxt.tag == qn('w:p'):
                if nxt.findall('.//' + qn('w:pict')):
                    has_img_after = True
                    break
                if nxt.findall('.//' + qn('w:t')):
                    break
        if has_img_after:
            to_remove.append(child)

    for p in to_remove:
        body.remove(p)


def clean_header_footer(doc, config=None):
    """Remove '(VARIABLE...)' markers from header/footer and update dates."""
    import datetime
    if config and config.get('fecha'):
        today = config['fecha']
    else:
        today = datetime.date.today().strftime('%d/%m/%Y')

    titulo = config.get('titulo', '') if config else ''
    codigo = config.get('codigo', '') if config else ''
    firma_fecha = config.get('firma', '') if config else ''

    for section in doc.sections:
        # Clean header
        for tbl in section.header.tables:
            _clean_table_variables(tbl, today)
            if titulo or codigo or config.get('fecha'):
                _apply_config_to_header_table(tbl, titulo, codigo, config.get('fecha', ''))
        for p in section.header.paragraphs:
            _clean_paragraph_variables(p, today)
        # Clean footer
        for tbl in section.footer.tables:
            _clean_table_variables(tbl, today)
            if firma_fecha:
                _apply_config_to_footer_table(tbl, firma_fecha)
        for p in section.footer.paragraphs:
            _clean_paragraph_variables(p, today)


def _clean_table_variables(tbl, today):
    """Clean VARIABLE markers from all cells in a table."""
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _clean_paragraph_variables(p, today)


def _clean_paragraph_variables(p, today):
    """Remove (VARIABLE...) from paragraph and update old dates.
    Works at full-paragraph level since text may be split across runs.
    """
    if not p.runs:
        return
    full = ''.join(r.text for r in p.runs)
    cleaned = re.sub(r'\s*\(VARIABLE[^)]*\)', '', full)
    cleaned = re.sub(r'\d{2}/\d{2}/\d{4}', today, cleaned)
    if cleaned != full:
        # Set first run to cleaned text, clear the rest
        p.runs[0].text = cleaned
        for r in p.runs[1:]:
            r.text = ''


def _clean_body_table_variables(tbl):
    """Clean VARIABLE markers from body table cells WITHOUT touching dates."""
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if not p.runs:
                    continue
                full = ''.join(r.text for r in p.runs)
                cleaned = re.sub(r'\s*\(VARIABLE[^)]*\)', '', full)
                if cleaned != full:
                    p.runs[0].text = cleaned
                    for r in p.runs[1:]:
                        r.text = ''


def add_page_break_before_table(doc, table):
    """Insert a page break paragraph immediately before a table in the document body.
    Skips if the preceding element is already a page break paragraph.
    """
    body = doc.element.body
    tbl_elem = table._tbl
    tbl_idx = list(body).index(tbl_elem)
    # Check if there's already a page break paragraph before this table
    if tbl_idx > 0:
        prev = list(body)[tbl_idx - 1]
        if prev.tag == qn('w:p'):
            for br in prev.iter(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    return  # Already has a page break before this table
    # Create a new paragraph with a page break
    p = etree.SubElement(body, qn('w:p'))
    r = etree.SubElement(p, qn('w:r'))
    br = etree.SubElement(r, qn('w:br'))
    br.set(qn('w:type'), 'page')
    # Move the paragraph before the table
    body.remove(p)
    body.insert(tbl_idx, p)


def _apply_config_to_header_table(tbl, titulo, codigo, fecha):
    """Apply titulo, codigo, and fecha to header table by cell position.

    Header table structure (4 rows x 3 visual cols):
      Row 0, Col 1: "Fecha: dd/mm/yyyy"
      Row 1, Col 0: "Unidad de Mediciones..."
      Row 2, Col 1: Codigo (+J01)  | Col 2: "pág. X"
      Row 3, Col 0: Titulo (INFORME DE PRUEBA...)
    """
    rows = tbl.rows
    if len(rows) < 4:
        return

    # Row 0, unique cell 1: fecha
    if fecha:
        cells_r0 = _get_unique_cells(rows[0])
        if len(cells_r0) >= 2:
            cell = cells_r0[1]
            for p in cell.paragraphs:
                if not p.runs:
                    continue
                full = ''.join(r.text for r in p.runs)
                new_text = re.sub(r'\d{2}/\d{2}/\d{4}', fecha, full)
                if new_text != full:
                    p.runs[0].text = new_text
                    for r in p.runs[1:]:
                        r.text = ''

    # Row 3, unique cell 0: titulo
    cells_r3 = _get_unique_cells(rows[3])
    if titulo and len(cells_r3) >= 1:
        set_cell_text(cells_r3[0], titulo)

    # Row 2, unique cell 1: codigo (preservar "pág. X")
    cells_r2 = _get_unique_cells(rows[2])
    if codigo and len(cells_r2) >= 2:
        cell = cells_r2[1]
        for p in cell.paragraphs:
            if not p.runs:
                continue
            full = ''.join(r.text for r in p.runs)
            # Replace codigo before "pág."
            new_text = re.sub(r'^[^\n]*?(?=pág\.)', codigo + ' ', full)
            if 'pág.' not in full:
                new_text = codigo
            if new_text != full:
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r.text = ''


def _apply_config_to_footer_table(tbl, firma_fecha):
    """Replace firma date in footer table (Row 2, Cell 0 = RIEDER firma)."""
    if len(tbl.rows) < 3:
        return
    cells = _get_unique_cells(tbl.rows[2])
    if not cells:
        return
    cell = cells[0]
    for p in cell.paragraphs:
        if not p.runs:
            continue
        full = ''.join(r.text for r in p.runs)
        new_text = re.sub(r'\d{2}/\d{2}/\d{4}', firma_fecha, full)
        if new_text == full and firma_fecha and 'Firma/Fecha:' in full:
            new_text = full.rstrip() + '\n' + firma_fecha
        if new_text != full:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''


def find_section(sections, name):
    """Find a section by name."""
    for s in sections:
        if s['name'] == name:
            return s
    return None


def update_header_table(tbl, data, param_col_idx, device_col_idx):
    """Update a Timing test header table with parameter values.

    Header tables have rows 2-7 with param values and device settings.
    Column indices vary by table (O1 has 7 cols, others have 5-6).
    """
    param_keys = [
        'Execution date/Time', 'Coil supply voltage', 'Measuring time',
        'Command pulse', 'Resolution', 'Frequency'
    ]
    device_keys = [
        'Device model / Serial-no', 'Transducer type',
        'Transducer on CBs phase-s', 'Transducer transfer functions',
        'Trigger source', 'BSG mode'
    ]
    params = data['params']
    device = data['device_settings']

    for i, (pk, dk) in enumerate(zip(param_keys, device_keys)):
        row_idx = i + 2
        if row_idx < len(tbl.rows):
            row = tbl.rows[row_idx]
            # Get non-merged cells
            seen = set()
            cells = []
            for c in row.cells:
                cid = id(c._tc)
                if cid not in seen:
                    seen.add(cid)
                    cells.append(c)
            # Set param value
            if param_col_idx < len(cells):
                set_cell_text(cells[param_col_idx], params.get(pk, ''))
            # Set device value
            if device_col_idx < len(cells):
                set_cell_text(cells[device_col_idx], device.get(dk, ''))


def update_data_rows(tbl, start_row, section_data):
    """Update data rows in a table starting at start_row.

    Each data row has: Designation | Min | Max | Value | Unit | Pass/Check
    We replace cols 1-5 (or all cols for the data portion).
    """
    if not section_data:
        return
    for di, data_row in enumerate(section_data):
        row_idx = start_row + di
        if row_idx >= len(tbl.rows):
            break
        row = tbl.rows[row_idx]
        seen = set()
        cells = []
        for c in row.cells:
            cid = id(c._tc)
            if cid not in seen:
                seen.add(cid)
                cells.append(c)
        # data_row is like: ['Open time(to)', '0', '0', '49.95', 'ms', 'Correcto']
        for ci in range(min(len(data_row), len(cells))):
            set_cell_text(cells[ci], data_row[ci])


def update_bounce_table(tbl, c_data):
    """Update bounce time table (table 8) with C test summary data."""
    bt = find_section(c_data['summary_sections'], 'Bounce time')
    if not bt:
        return
    # Rows 2-4: Bounce time-A, B, C
    for di, data_row in enumerate(bt['data']):
        row_idx = 2 + di
        if row_idx >= len(tbl.rows):
            break
        row = tbl.rows[row_idx]
        seen = set()
        cells = []
        for c in row.cells:
            cid = id(c._tc)
            if cid not in seen:
                seen.add(cid)
                cells.append(c)
        # data_row from summary: ['Bounce time-A', '0', '0', '1.925', 'ms', ...]
        # Template cols: Designation | C1 | Unit
        if len(cells) >= 3:
            set_cell_text(cells[0], data_row[0])  # Designation
            set_cell_text(cells[1], data_row[3] if len(data_row) > 3 else '')  # Value (C1)
            set_cell_text(cells[2], data_row[4] if len(data_row) > 4 else 'ms')  # Unit


def update_coil_currents_table(tbl, co_data):
    """Update coil currents table (table 14) with CO test data."""
    coils = [s for s in co_data['sections'] if s['name'] == 'Coil currents']
    if not coils:
        return
    # First coil section (Closing): rows 2-3
    if len(coils) >= 1 and coils[0]['data']:
        for di, dr in enumerate(coils[0]['data']):
            ri = 2 + di
            if ri >= len(tbl.rows):
                break
            row = tbl.rows[ri]
            seen = set()
            cells = []
            for c in row.cells:
                cid = id(c._tc)
                if cid not in seen:
                    seen.add(cid)
                    cells.append(c)
            for ci in range(min(len(dr), len(cells))):
                set_cell_text(cells[ci], dr[ci])
    # Second coil section (Opening): rows 7-8
    if len(coils) >= 2 and coils[1]['data']:
        for di, dr in enumerate(coils[1]['data']):
            ri = 7 + di
            if ri >= len(tbl.rows):
                break
            row = tbl.rows[ri]
            seen = set()
            cells = []
            for c in row.cells:
                cid = id(c._tc)
                if cid not in seen:
                    seen.add(cid)
                    cells.append(c)
            for ci in range(min(len(dr), len(cells))):
                set_cell_text(cells[ci], dr[ci])


# =============================================================================
# IMAGE REPLACEMENT (ZIP level)
# =============================================================================

def get_body_image_info(docx_path):
    """Get VML image info: list of (rId, shape_height_pt) in document order."""
    with zipfile.ZipFile(docx_path, 'r') as z:
        doc_xml = z.read('word/document.xml')
    tree = etree.fromstring(doc_xml)
    results = []
    seen = set()
    for shape in tree.iter('{' + VML_NS + '}shape'):
        imgdata = shape.find('{' + VML_NS + '}imagedata')
        if imgdata is None:
            continue
        rId = imgdata.get(qn('r:id'))
        if not rId or rId in seen:
            continue
        seen.add(rId)
        # Parse height from style attribute
        style = shape.get('style', '')
        height_pt = 0
        m = re.search(r'height:\s*([\d.]+)(pt|in)', style)
        if m:
            val = float(m.group(1))
            height_pt = val * 72 if m.group(2) == 'in' else val
        results.append((rId, height_pt))
    return results


def get_rel_targets(docx_path):
    """Get rId -> target path mapping from document.xml.rels."""
    with zipfile.ZipFile(docx_path, 'r') as z:
        rels_xml = z.read('word/_rels/document.xml.rels')
    tree = etree.fromstring(rels_xml)
    mapping = {}
    for rel in tree:
        mapping[rel.get('Id')] = rel.get('Target')
    return mapping


def replace_images_in_zip(docx_path, output_path, image_replacements):
    """Replace media files in the .docx ZIP.

    image_replacements: dict of {media_path: new_bytes}
    e.g. {'word/media/image1.jpeg': b'...'}
    """
    temp_path = output_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in image_replacements:
                    zout.writestr(item, image_replacements[item.filename])
                else:
                    zout.writestr(item, zin.read(item.filename))
    # Move temp to output
    if os.path.exists(output_path):
        os.remove(output_path)
    os.rename(temp_path, output_path)


# =============================================================================
# APPLY USER-EDITED FIELDS TO TEMPLATE
# =============================================================================

def _get_unique_cells(row):
    """Get unique (non-merged-duplicate) cells from a table row."""
    seen = set()
    cells = []
    for c in row.cells:
        cid = id(c._tc)
        if cid not in seen:
            seen.add(cid)
            cells.append(c)
    return cells


def apply_pass_overrides(doc, fields):
    """Apply pass/check override values from HTML preview selects to template tables.

    Field keys have format: pass_{slot}_{sectionName}_{rowIndex}
    Maps each to the correct template table/row and updates the last unique cell (Pass/Check column).
    """
    if not fields:
        return

    # Map: (slot, section_name) -> (table_idx, start_row)
    # start_row is the first data row in the template table for that section
    SECTION_MAP = {
        ('O1', 'Main contacts timing'):    (5, 2),
        ('O1', 'Opening time'):            (5, 7),
        ('O1', 'Opening coil result 1'):   (5, 14),  # rows 14-15 in T5
        ('O1', 'Opening coil result 1 pt2'): (6, 2),  # T6 row 2+
        ('C', 'Main contacts timing'):     (9, 2),
        ('C', 'Closing time'):             (9, 7),
        ('OCO', 'Main contacts timing'):   (11, 2),
        ('OCO', 'Opening time'):           (11, 12),
        ('CO', 'Main contacts timing'):    (13, 2),
        ('CO', 'Closing time'):            (13, 9),
        ('CO', 'Coil currents'):           (14, 2),  # Closing coils (rows 2-3)
        ('CO_coil1', 'Coil currents'):    (14, 7),  # Opening coils (rows 7-8)
        ('O2', 'Main contacts timing'):    (16, 2),
        ('O2', 'Opening time'):            (16, 7),
        ('O2', 'Opening coil result 1'):   (16, 14),  # rows 14-15 in T16
        ('O2', 'Opening coil result 1 pt2'): (17, 2),  # T17 row 2+
    }

    for key, value in fields.items():
        if not key.startswith('pass_'):
            continue
        # Parse: pass_{slot}_{sectionName}_{rowIndex}
        parts = key[5:]  # remove 'pass_'
        # Find the last underscore to get rowIndex
        last_under = parts.rfind('_')
        if last_under < 0:
            continue
        try:
            row_idx = int(parts[last_under + 1:])
        except ValueError:
            continue
        rest = parts[:last_under]
        # Find slot: first segment before underscore
        first_under = rest.find('_')
        if first_under < 0:
            continue
        slot = rest[:first_under]
        section_name = rest[first_under + 1:]

        lookup = (slot, section_name)
        if lookup not in SECTION_MAP:
            continue

        tbl_idx, start_row = SECTION_MAP[lookup]
        actual_row = start_row + row_idx
        if tbl_idx < len(doc.tables):
            tbl = doc.tables[tbl_idx]
            if actual_row < len(tbl.rows):
                cells = _get_unique_cells(tbl.rows[actual_row])
                if cells:
                    set_cell_text(cells[-1], value)


def apply_fields(doc, fields):
    """Apply user-edited field values from HTML preview to template tables."""
    if not fields:
        return

    # Field mapping: field_name -> (table_idx, row_idx, unique_cell_idx, label_prefix)
    # label_prefix: if set, cell text becomes "prefix + value"; if None, entire cell = value
    FIELD_MAP = {
        'obra':                (1, 2, 0, 'OBRA/LUGAR: '),
        'fabricante':          (1, 4, 0, 'FABRICANTE: '),
        'tension':             (1, 4, 1, 'TENSION ASIGNADA NORMAL: '),
        'tipo':                (1, 5, 0, 'TIPO: '),
        'corriente':           (1, 5, 1, 'CORRIENTE ASIGNADA NORMAL: '),
        'anio':                (1, 6, 0, u'A\u00d1O FABRICACION: '),
        'cortocircuito':       (1, 6, 1, 'CORRIENTE DE CORTOCIRCUITO: '),
        'serial':              (1, 7, 0, 'NUMERO DE SERIE: '),
        'frecuencia':          (1, 7, 1, 'FRECUENCIA NOMINAL: '),
        'equipo_modelo':       (2, 3, 2, None),
        'equipo_calibracion':  (2, 3, 3, None),
        'observaciones':       (3, 1, 0, None),
    }

    for field_name, (tbl_idx, row_idx, cell_idx, prefix) in FIELD_MAP.items():
        value = fields.get(field_name)
        if value is None:
            continue
        tbl = doc.tables[tbl_idx]
        if row_idx >= len(tbl.rows):
            continue
        cells = _get_unique_cells(tbl.rows[row_idx])
        if cell_idx >= len(cells):
            continue
        cell = cells[cell_idx]
        new_text = (prefix + value) if prefix else value
        set_cell_text(cell, new_text)

    # NOTE: pass/check overrides are applied AFTER update_data_rows in generate()

    # Apply verification visual states via Wingdings checkbox symbols
    # Each cell has 2 paragraphs with w:sym elements: p0="Buen estado", p1="Mal estado"
    # We toggle the w:char attribute: F0FE=☑ checked, F06F=☐ unchecked
    CHECKED = 'F0FE'
    UNCHECKED = 'F06F'
    verificacion = fields.get('verificacion')
    if verificacion and isinstance(verificacion, list):
        for i, item in enumerate(verificacion):
            row_idx = 10 + i  # Table 1 rows 10-14
            if row_idx >= len(doc.tables[1].rows):
                break
            cells = _get_unique_cells(doc.tables[1].rows[row_idx])
            # cells: [label, Fase A, Fase B, Fase C]
            for pi, phase in enumerate(['a', 'b', 'c']):
                ci = 1 + pi
                if ci < len(cells):
                    val = item.get(phase, 'buen')
                    cell = cells[ci]
                    for p_idx, p in enumerate(cell.paragraphs):
                        sym = p._element.find('.//' + qn('w:sym'))
                        if sym is not None:
                            if (p_idx == 0 and val == 'buen') or (p_idx == 1 and val == 'mal'):
                                sym.set(qn('w:char'), CHECKED)
                            else:
                                sym.set(qn('w:char'), UNCHECKED)


# =============================================================================
# MAIN
# =============================================================================

def generate(input_files, template_path, output_path, config=None, fields=None):
    """Generate report from input files using template.

    Args:
        input_files: dict {'O1': path, 'C': path, 'OCO': path, 'CO': path, 'O2': path}
        template_path: path to template.docx
        output_path: path for output .docx
        config: optional dict with {titulo, codigo, fecha, firma} from UI
        fields: optional dict with editable field values from HTML preview
    """
    if config is None:
        config = {}
    if fields is None:
        fields = {}

    print(f"  Config received: {list(config.keys()) if config else 'EMPTY'}")
    print(f"  Fields received: {len(fields)} fields")
    if config:
        print(f"    titulo={config.get('titulo', '')[:40]}")
        print(f"    codigo={config.get('codigo', '')}")
        print(f"    fecha={config.get('fecha', '')}")
        print(f"    firma={config.get('firma', '')}")

    # Step 1: Parse all input files
    print("Step 1: Parsing input files...")
    all_data = {}
    for slot, filepath in input_files.items():
        all_data[slot] = parse_input_file(filepath, slot)
        print(f"    {slot}: {len(all_data[slot]['sections'])} sections, {len(all_data[slot]['images'])} images")

    # Step 2: Open template and modify table data
    print("\nStep 2: Modifying template tables...")
    doc = Document(template_path)

    # Center all tables and VML images horizontally
    print("  Centering all tables and images...")
    center_all_content(doc)

    # Clean (VARIABLE...) markers from header, footer, and all body tables
    print("  Cleaning VARIABLE markers...")
    clean_header_footer(doc, config)
    for tbl in doc.tables:
        _clean_body_table_variables(tbl)

    # -- Front matter (Tables 0-3): apply user-edited fields --
    o1 = all_data['O1']
    # If fields not provided from UI, at minimum use O1 data for tipo/serial
    if not fields.get('tipo'):
        fields['tipo'] = o1['header'].get('type', '8DA10')
    if not fields.get('serial'):
        fields['serial'] = o1['header'].get('serial', '')
    print("  Applying user-edited fields...")
    apply_fields(doc, fields)

    # -- O1: Table 4 (header), Table 5 (data), Table 6 (coil part 2) --
    print("  Updating O1...")
    update_header_table(doc.tables[4], o1, param_col_idx=1, device_col_idx=3)
    o1_main = find_section(o1['sections'], 'Main contacts timing')
    o1_opening = find_section(o1['sections'], 'Opening time')
    o1_coil = find_section(o1['sections'], 'Opening coil result 1')
    if o1_main:
        update_data_rows(doc.tables[5], 2, o1_main['data'])
    if o1_opening:
        update_data_rows(doc.tables[5], 7, o1_opening['data'])
    if o1_coil and len(o1_coil['data']) >= 2:
        update_data_rows(doc.tables[5], 14, o1_coil['data'][:2])
    if o1_coil and len(o1_coil['data']) > 2:
        update_data_rows(doc.tables[6], 2, o1_coil['data'][2:])

    # -- C: Table 7 (header), Table 8 (bounce), Table 9 (data) --
    print("  Updating C...")
    c_data = all_data['C']
    update_header_table(doc.tables[7], c_data, param_col_idx=1, device_col_idx=3)
    update_bounce_table(doc.tables[8], c_data)
    c_main = find_section(c_data['sections'], 'Main contacts timing')
    c_closing = find_section(c_data['sections'], 'Closing time')
    if c_main:
        update_data_rows(doc.tables[9], 2, c_main['data'])
    if c_closing:
        update_data_rows(doc.tables[9], 7, c_closing['data'])

    # -- OCO: Table 10 (header), Table 11 (data) --
    print("  Updating OCO...")
    oco = all_data['OCO']
    update_header_table(doc.tables[10], oco, param_col_idx=1, device_col_idx=3)
    oco_main = find_section(oco['sections'], 'Main contacts timing')
    oco_opening = find_section(oco['sections'], 'Opening time')
    if oco_main:
        update_data_rows(doc.tables[11], 2, oco_main['data'])
    if oco_opening:
        update_data_rows(doc.tables[11], 12, oco_opening['data'])

    # -- CO: Table 12 (header), Table 13 (data), Table 14 (coils) --
    print("  Updating CO...")
    co = all_data['CO']
    update_header_table(doc.tables[12], co, param_col_idx=1, device_col_idx=3)
    co_main = find_section(co['sections'], 'Main contacts timing')
    co_closing = find_section(co['sections'], 'Closing time')
    if co_main:
        update_data_rows(doc.tables[13], 2, co_main['data'])
    if co_closing:
        update_data_rows(doc.tables[13], 9, co_closing['data'])
    update_coil_currents_table(doc.tables[14], co)

    # -- O2: Table 15 (header), Table 16 (data), Table 17 (coil part 2) --
    print("  Updating O2...")
    o2 = all_data['O2']
    update_header_table(doc.tables[15], o2, param_col_idx=1, device_col_idx=3)
    o2_main = find_section(o2['sections'], 'Main contacts timing')
    o2_opening = find_section(o2['sections'], 'Opening time')
    o2_coil = find_section(o2['sections'], 'Opening coil result 1')
    if o2_main:
        update_data_rows(doc.tables[16], 2, o2_main['data'])
    if o2_opening:
        update_data_rows(doc.tables[16], 7, o2_opening['data'])
    if o2_coil and len(o2_coil['data']) >= 2:
        update_data_rows(doc.tables[16], 14, o2_coil['data'][:2])
    if o2_coil and len(o2_coil['data']) > 2:
        update_data_rows(doc.tables[17], 2, o2_coil['data'][2:])

    # Apply pass/check overrides AFTER all update_data_rows (so they don't get overwritten)
    print("  Applying pass/check overrides...")
    apply_pass_overrides(doc, fields)

    # Insert page break in Table 1 before DESCRIPCION GENERAL (row 3)
    # so Page 1 = LISTADO + DATOS GENERALES (cliente, obra)
    # and Page 2 = DESCRIPCION + VERIFICACION + TIEMPOS + OBSERVACIONES
    print("  Inserting page break before DESCRIPCION GENERAL...")
    desc_row = doc.tables[1].rows[3]
    desc_cell = desc_row.cells[0]
    if desc_cell.paragraphs:
        p_elem = desc_cell.paragraphs[0]._p
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p_elem, qn('w:pPr'))
            p_elem.insert(0, pPr)
        pgBr = etree.SubElement(pPr, qn('w:pageBreakBefore'))

    # Move Bounce time table (T8) after C images so images stay together on one page
    print("  Moving Bounce table after C images...")
    body = doc.element.body
    t8_elem = doc.tables[8]._tbl
    t9_elem = doc.tables[9]._tbl
    body.remove(t8_elem)
    body.insert(list(body).index(t9_elem), t8_elem)

    # Force Bounce table (T8) to start on next page (avoid splitting across pages)
    print("  Adding page break before Bounce table...")
    bounce_cell = doc.tables[8].rows[0].cells[0]
    if bounce_cell.paragraphs:
        bp = bounce_cell.paragraphs[0]._p
        bpPr = bp.find(qn('w:pPr'))
        if bpPr is None:
            bpPr = etree.SubElement(bp, qn('w:pPr'))
            bp.insert(0, bpPr)
        etree.SubElement(bpPr, qn('w:pageBreakBefore'))

    # Insert page breaks before each timing test header table
    print("  Inserting page breaks before timing tests...")
    for tbl_idx in [4, 7, 10, 12, 15]:
        add_page_break_before_table(doc, doc.tables[tbl_idx])

    # Save intermediate file
    intermediate = output_path + '.intermediate.docx'
    doc.save(intermediate)
    print(f"  Saved intermediate: {intermediate}")

    # Step 3: Replace images at ZIP level
    print("\nStep 3: Replacing images...")
    img_info = get_body_image_info(intermediate)
    rels = get_rel_targets(intermediate)
    print(f"  Found {len(img_info)} body image references")

    # Collect input images per slot: each has [strip (short), graph (tall)]
    # From input .docx: image00001=strip (242px h), image00002=graph (722px h)
    slot_images = {}
    for slot in ['O1', 'C', 'OCO', 'CO', 'O2']:
        imgs = all_data[slot]['images']
        if len(imgs) >= 2:
            # Determine which is strip vs graph by checking JPEG dimensions
            from PIL import Image
            dims = []
            for img_bytes in imgs:
                im = Image.open(io.BytesIO(img_bytes))
                dims.append(im.size)  # (w, h)
                im.close()
            # Strip is shorter (h < w/3), graph is taller
            if dims[0][1] < dims[1][1]:
                slot_images[slot] = {'strip': imgs[0], 'graph': imgs[1]}
            else:
                slot_images[slot] = {'strip': imgs[1], 'graph': imgs[0]}
            print(f"  {slot}: strip={dims[0][0]}x{dims[0][1]}, graph={dims[1][0]}x{dims[1][1]}")

    # Template shapes come in pairs per slot: strip (~86pt) and graph (~259pt)
    # Match by shape height: < 150pt = strip, >= 150pt = graph
    STRIP_THRESHOLD = 150  # pt
    slots_order = ['O1', 'C', 'OCO', 'CO', 'O2']
    slot_idx = 0
    pair_count = 0  # count shapes assigned per slot (2 per slot)

    replacements = {}
    for rId, height_pt in img_info:
        slot = slots_order[slot_idx]
        is_strip = height_pt < STRIP_THRESHOLD
        img_type = 'strip' if is_strip else 'graph'
        target = rels.get(rId, '')
        if target and slot in slot_images:
            media_path = 'word/' + target if not target.startswith('/') else target.lstrip('/')
            replacements[media_path] = slot_images[slot][img_type]
            print(f"  {media_path} (rId={rId}, h={height_pt:.1f}pt) <- {slot} {img_type}")
        pair_count += 1
        if pair_count >= 2:
            pair_count = 0
            slot_idx += 1
            if slot_idx >= len(slots_order):
                break

    replace_images_in_zip(intermediate, output_path, replacements)

    if os.path.exists(intermediate):
        os.remove(intermediate)

    print(f"\nDONE - Report generated: {output_path}")


def main():
    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERROR: Template not found: {TEMPLATE_PATH}")
        print("Run convert_template.py first to convert .doc to .docx")
        sys.exit(1)
    generate(INPUT_FILES, TEMPLATE_PATH, OUTPUT_PATH)


def slot_for_idx(i):
    """Map image index to slot name for logging."""
    slots = ['O1', 'O1', 'C', 'C', 'OCO', 'OCO', 'CO', 'CO', 'O2', 'O2']
    return slots[i] if i < len(slots) else '?'


if __name__ == '__main__':
    main()
